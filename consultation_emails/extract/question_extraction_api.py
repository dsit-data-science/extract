"""
Classes:
    LangchainExtractor: Extracts responses from a PDF or DOCX document using a language model and a set of questions.

    ResponseUploader: Handle the responses from the model being uploaded to database


Usage example:

    .. code-block:: python

        pdf_filepath = "your_pdf_responses_filepath.pdf"  # or .docx
        csv_filepath = "your_csv_questions_filepath.csv"

        from consultation_emails.extract.question_extraction_api import LangchainExtractor
        from consultation_emails.database.fetch_data import StorageHandler, DynamoDbHandler
        from consultation_emails.extract.question_extraction_api import ResponseUploader

        storage_handler = StorageHandler()
        dynamo_db_handler = DynamoDbHandler()
        langchain_extractor = LangchainExtractor(use_environment_variables=False)

        human_prompt_template = storage_handler.load_jinja("prompt_templates/approved/prompt_template_human.jinja")
        system_prompt_template = storage_handler.load_jinja("prompt_templates/approved/prompt_template_system.jinja")
        responses = langchain_extractor.extract_responses(
            document_filepath=pdf_filepath,
            csv_filepath=csv_filepath,
            system_prompt_template=system_prompt_template,
            human_prompt_template=human_prompt_template,
        )

        response_uploader = ResponseUploader(dynamo_db_handler=dynamo_db_handler)
        doc_text = langchain_extractor.load_document_text(document_filepath=pdf_filepath)
        doc_id = response_uploader.create_doc_id(doc_text)
        response_uploader.upload_responses(doc_id=doc_id, responses=responses.model_dump())

"""  # noqa: E501

import hashlib
import os
from datetime import datetime
from decimal import Decimal
from io import BytesIO

import fitz
from docx import Document as DocxDocument

from langchain.prompts import ChatPromptTemplate
from langchain_core.documents import Document
from langchain_core.exceptions import OutputParserException
from langchain_core.output_parsers import PydanticOutputParser
from langchain_openai import AzureChatOpenAI
from tenacity import RetryCallState, Retrying, retry_if_exception_type, stop_after_attempt

from consultation_emails.database.fetch_data import DynamoDbHandler, SecretsManagerHandler, SsmHandler, StorageHandler
from consultation_emails.extract.text_handler import PlaintextHolder
from consultation_emails.schemas.dynamo_db import ManualIdSchema, ModelOutputsSchema, QuestionSchema, ResponseIdSchema, ReviewIdSchema
from consultation_emails.schemas.model_responses import QuestionAnswer, QuestionAnswers


class LangchainExtractor:
    def __init__(
        self,
        use_environment_variables: bool = False,
    ):
        """
        Initializes the LangchainExtractor class.

        Args:
            use_environment_variables (bool, optional): Flag indicating whether to use environment variables for configuration.
                Defaults to False.
            api_version (str, optional): The API version to use. Defaults to "2023-03-15-preview".
            azure_deployment (str, optional): The Azure deployment to use. Defaults to None.
            api_key (str, optional): The API key to use. Defaults to None.

        Class Attributes:
            _azure_client (AzureChatOpenAI): The Azure OpenAI client to use for extracting responses.

        Environment variables used:
            - AZURE_OPENAI_API_VERSION: The API version to use.
            - AZURE_OPENAI_CHAT_DEPLOYMENT_NAME: The Azure deployment to use.
            - AZURE_OPENAI_API_KEY: The API key to use.

        """
        self.env_vars = use_environment_variables
        self._secrets_manager_handler = SecretsManagerHandler()

        secret_id = os.environ.get("MODEL_SECRETS_ID", "extract-secrets")
        self._secrets = self._secrets_manager_handler.get_secret_value(secret_id=secret_id)

        self._ssm_handler = SsmHandler()

        self._azure_client = self._get_model_client()

        self._storage_handler = StorageHandler()

    def _get_model_client(self) -> AzureChatOpenAI:
        """
        Returns an instance of the AzureChatOpenAI client.

        Args:
            api_version (str, optional): The API version to use. Defaults to "2023-03-15-preview".
            azure_deployment (str, optional): The Azure deployment to use. Defaults to None.
            api_key (str, optional): The API key to use. Defaults to None.
        Returns:
            AzureChatOpenAI: An instance of the AzureChatOpenAI client configured with the specified parameters.

        Environment variables used:
            - AZURE_OPENAI_API_VERSION: The API version to use.
            - AZURE_OPENAI_CHAT_DEPLOYMENT_NAME: The Azure deployment to use.
            - AZURE_OPENAI_API_KEY: The API key to use.

        """
        # Get model configuration
        temperature=self._ssm_handler.get_parameter("extract-model-temperature")

        if self.env_vars:
            model = AzureChatOpenAI(
                azure_endpoint=os.environ["AZURE_OPENAI_ENDPOINT"],
                openai_api_version=os.environ["AZURE_OPENAI_API_VERSION"],
                azure_deployment=os.environ["AZURE_OPENAI_CHAT_DEPLOYMENT_NAME"],
                api_key=os.environ["AZURE_OPENAI_API_KEY"],
                temperature=float(temperature),
            )

        else:
            openai_api_version = self._secrets["AZURE_OPENAI_API_VERSION"]
            azure_deployment = self._secrets["AZURE_OPENAI_CHAT_DEPLOYMENT_NAME"]
            api_key = self._secrets["AZURE_OPENAI_API_KEY"]
            azure_openai_endpoint = self._secrets["AZURE_OPENAI_ENDPOINT"]

            model = AzureChatOpenAI(
                azure_endpoint=azure_openai_endpoint,
                openai_api_version=openai_api_version,
                azure_deployment=azure_deployment,
                api_key=api_key,
                temperature=float(temperature),
            )
            
        return model

    @staticmethod
    def before_retry(retry_state: RetryCallState):
        print(f"Trying chain. Attempt number: {retry_state.attempt_number}")

    def extract_responses(
        self, document_filepath: str, csv_filepath: str, system_prompt_template: str, human_prompt_template: str
    ) -> QuestionAnswers:
        """
        Extracts responses from a PDF or DOCX document using a language model and a set of questions.
            Variables used in the system message template:
                - {{ json_schema_instructions }}: The JSON schema instructions for the output format.
                - {{ questions }}: A list of JSON objects containing 'question_label' and 'question_text' fields.


        Args:
            document_filepath (str): The file path to the PDF or DOCX document.
            csv_filepath (str): The file path to the CSV file containing the questions.
            system_prompt_template (str): The Jinja template for the system message.
            human_prompt_template (str): The Jinja template for the human message.

        Returns:
            dict: The extracted responses in JSON format.

        """

        response_text = self.load_document_text(document_filepath)

        questions = self._storage_handler.load_csv(csv_filepath)

        prompt = ChatPromptTemplate([("system", system_prompt_template), ("human", human_prompt_template)], template_format="jinja2")

        pydantic_parser = PydanticOutputParser(pydantic_object=QuestionAnswers)

        chain = prompt | self._azure_client | pydantic_parser 

        input_data = {
            "json_schema_instructions": pydantic_parser.get_format_instructions(),
            "response_text": response_text,
            "questions": questions,
        }

        # Call the chain, retrying if an OutputParserException is raised
        retry_wrapper = Retrying(retry=retry_if_exception_type(OutputParserException), stop=stop_after_attempt(2), before=self.before_retry)
        question_answers = retry_wrapper(chain.invoke, input=input_data)
        question_answers.count_retries = retry_wrapper.statistics.get("attempt_number", 1) - 1

        plaintext_holder = PlaintextHolder(response_text)

        # Replace any imperfect model extractions with their closest extraction using Jaccard similarity
        for question_answer in question_answers.question_answers:
            matched_extractions = []
            similarity_scores = []
            for extracted_string in question_answer.extracted_text:
                exact_search_results = plaintext_holder.exact_text_search(extracted_string, max_results=1)
                if exact_search_results:
                    matched_text = plaintext_holder.convert_tokens_to_string(exact_search_results[0]["tokens"])
                    matched_extractions.append(matched_text)
                    similarity_scores.append(Decimal("1.0"))
                else:
                    search_result = plaintext_holder.jaccard_similarity_search(
                        text=extracted_string,
                    )[0]
                    closest_matched_text = plaintext_holder.convert_tokens_to_string(search_result["tokens"])
                    matched_extractions.append(closest_matched_text)
                    similarity_scores.append(Decimal(str(search_result["score"])))

            question_answer.original_extracted_text = question_answer.extracted_text
            question_answer.extracted_text = matched_extractions
            question_answer.extracted_text_similarity = similarity_scores

        return question_answers

    def load_document_text(self, document_filepath: str):
        """
        Loads the text content of a document (PDF, DOCX).

        Args:
            document_filepath (str): The file path to the document.
        Returns:
            str: The text content of the document.

        """
        if document_filepath.lower().endswith(".pdf"):
            return self._load_pdf_text(key=document_filepath)
        elif document_filepath.lower().endswith(".docx"):
            return self._load_docx_text(key=document_filepath)
        else:
            raise ValueError("Unsupported file format. Only PDF and DOCX are supported.")

    def _load_pdf_text(self, key: str, bucket: str = None) -> str:
        """
        Loads the text content of a PDF file from S3 storage.

        Args:
            key (str): The S3 key (file path) to the PDF document.
            bucket (str, optional): The S3 bucket name. If not provided, uses the USER_UPLOAD_S3_BUCKET environment variable.

        Returns:
            str: The concatenated text content of the PDF document, with each page separated by double newlines.

        """
        if not bucket:
            bucket = os.environ["USER_UPLOAD_S3_BUCKET"]

        response = self._storage_handler.get_object(
            key=key, 
            bucket=bucket,
        )
        response_bytes = response["Body"].read()
        reader = fitz.open(stream=BytesIO(response_bytes))

        doc = []
        for page_num in range(reader.page_count):
            page = reader.load_page(page_num)
            doc.append(Document(page_content=page.get_text()))

        return self._join_document_pages(doc)

    def _load_docx_text(self, key: str, bucket: str = None):
        """
        Loads the text content of a DOCX file from S3 storage.

        Args:
            key (str): The S3 key (file path) to the DOCX document.
            bucket (str, optional): The S3 bucket name. If not provided, uses the USER_UPLOAD_S3_BUCKET environment variable.

        Returns:
            str: The concatenated text content of the DOCX document, with each paragraph separated by double newlines.

        """
        if not bucket:
            bucket = os.environ["USER_UPLOAD_S3_BUCKET"]

        response = self._storage_handler.get_object(
            key=key, 
            bucket=bucket,
        )
        response_bytes = response["Body"].read()
        stream = BytesIO(response_bytes)

        docx = DocxDocument(stream)

        doc = []
        for para in docx.paragraphs:
            text = para.text.strip()
            if text:
                doc.append(Document(page_content=text))

        return self._join_document_pages(doc)

    def _join_document_pages(self, docs: list[Document]) -> str:
        """
        Joins the page content of a list of Document objects into a single string.

        Args:
            docs (list[langchain_core.documents.Document]): A list of langchain Document objects with page_content fields.
        Returns:
            str: A single string with the page content of each Document object separated by double newlines.

        """
        return "\n\n".join(doc.page_content for doc in docs)


class ResponseUploader:
    """
    Handle the responses from the model being uploaded to database
        Args:
            dynamo_db_handler: An instance of the DynamoDbHandler for uploading to DynamoDB

    """

    def __init__(self, dynamo_db_handler: DynamoDbHandler):
        self._dynamo_db_handler = dynamo_db_handler

    def create_doc_id(self, response_text: str):
        """
        Generates a SHA-256 hash ID for the given response text.

        Args:
            response_text (str): The text to be hashed.
        Returns:
            str: The SHA-256 hash of the response text in hexadecimal format.

        """
        response_bytes = response_text.encode("utf-8")
        hash_object = hashlib.sha256(response_bytes)
        return hash_object.hexdigest()

    def _validate_responses(self, doc_id: str, responses: dict, metadata: dict = None):
        return ModelOutputsSchema(
            id=doc_id,
            processed_datetime=int(datetime.now().timestamp()),
            item=responses.get("question_answers", {}),
            metadata=metadata,
            count_retries=responses.get("count_retries", Decimal("0")),
        )

    def upload_responses(
        self,
        doc_id: str,
        responses: dict,
        metadata: dict = None,
        responses_table_name: str = None,
    ):
        """
        Uploads the responses to the database after validation.

        Args:
            doc_id (str): The ID of the document.
            responses (dict): A dictionary containing the responses to be uploaded.
            metadata (dict, optional): Additional metadata for the responses. Defaults to None.
            responses_table_name (str, optional): The name of the table to upload the responses to.
                Defaults to the environment variable RESPONSES_TABLE_NAME.
        Returns:
            None

        """
        if responses_table_name is None:
            responses_table_name = os.environ["RESPONSES_TABLE_NAME"]
        if metadata is None:
            metadata = {}

        validated_responses = self._validate_responses(doc_id=doc_id, responses=responses, metadata=metadata)

        self._dynamo_db_handler.upload_data(data=validated_responses.__dict__, table_name=responses_table_name)

    def create_question_schema(
        self,
        question_answer: QuestionAnswer,
        doc_id: str,
    ) -> QuestionSchema:
        question_schema = QuestionSchema()
        question_schema.response_id = doc_id
        question_schema.timestamp = int(datetime.now().timestamp())
        question_schema.question_label = question_answer.question_label
        question_schema.question_text = question_answer.question_text
        question_schema.extracted_text = question_answer.original_extracted_text
        question_schema.exact_extracted_text = question_answer.extracted_text
        question_schema.is_exact_match = [similarity_score == Decimal("1.0") for similarity_score in question_answer.extracted_text_similarity]
        question_schema.jaccard_similarity = question_answer.extracted_text_similarity
        return question_schema

    def create_response_schema(
        self,
        doc_id: str,
        responses: QuestionAnswers,
        prompt_template: list[str],
        plain_text: str,
    ) -> ResponseIdSchema:
        response_schema = ResponseIdSchema()
        response_schema.response_id = doc_id
        response_schema.timestamp = int(datetime.now().timestamp())
        response_schema.retries = responses.count_retries
        response_schema.prompt_template = prompt_template
        response_schema.plain_text = plain_text
        response_schema.model = [os.getenv("AZURE_OPENAI_CHAT_DEPLOYMENT_NAME", ""), os.getenv("AZURE_OPENAI_API_VERSION", "")]
        return response_schema

    def create_review_schema(
        self,
        question_answer: QuestionAnswer,
        doc_id: str,
    ) -> ReviewIdSchema:
        review_schema = ReviewIdSchema()
        review_schema.response_id = doc_id
        review_schema.question_label = question_answer.question_label
        review_schema.user_id = ""  # Add user ID here once integration to Cognito is fixed
        review_schema.timestamp = None
        review_schema.reviewed_text = []
        review_schema.status = "None"

        return review_schema

    def create_manual_schema(
        self,
        question_answer: QuestionAnswer,
        doc_id: str,
    ) -> ManualIdSchema:
        """
        Creates a manual extraction schema for the given question answer and document ID.
        Args:
            question_answer (QuestionAnswer): The question answer object containing the extraction data.
            doc_id (str): The document ID of the response.
        Returns:
            ManualIdSchema: A new ManualIdSchema object with the same structure as the input, but with 'empty' answers removed.
        """
        manual_schema = ManualIdSchema()
        manual_schema.response_id = doc_id
        manual_schema.question_label = question_answer.question_label
        manual_schema.user_id = ""  # Add user ID here once integration to Cognito is fixed
        manual_schema.status = "Original"

        return manual_schema
